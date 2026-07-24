"""
Visual styles and element factories for the pipeline report.

All colors, fonts, and sizes used by the renderer are defined here so that
visual identity can be tuned in one place. Helper functions wrap the python-docx
internals required to produce filled banner tables, callout boxes, paragraph
borders, and page breaks — operations that python-docx exposes only through
its lower-level OOXML API.
"""

from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


# --- Color palette --------------------------------------------------------
# Hex strings (no leading "#") — used directly in OOXML attributes.

NAVY = "1F4E79"           # Primary heading and section banner color
ORANGE = "ED7D31"         # Cover-page subtitle color
LIGHT_BLUE = "D9E2F3"     # Callout box fill (metadata strips, status panels)
WHITE = "FFFFFF"          # Used for text on dark-navy banners
BORDER_GREY = "BFBFBF"    # Subtle table cell borders


# --- Font + size constants ------------------------------------------------

FONT = "Arial"

# Sizes are in points. python-docx wraps with Pt() at use site.
SIZE_BODY = 11
SIZE_COVER_TITLE = 28
SIZE_COVER_SUBTITLE = 22
SIZE_STEP_HEADING = 22
SIZE_SECTION_HEADING = 16
SIZE_SUBSECTION_HEADING = 13
SIZE_BANNER = 13
SIZE_HEADER_FOOTER = 9


# --- Helper: RGB tuple from hex -------------------------------------------

def _rgb(hex_color):
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


# --- Helper: apply font to a run ------------------------------------------

def style_run(run, *, size=SIZE_BODY, bold=False, italic=False, color=None, font_name=FONT):
    """Apply font name, size, weight, and color to a run in one call."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = _rgb(color)


# --- Helper: paragraph borders (horizontal rules) -------------------------

# OOXML schema requires pPr children to appear in a specific order. pBdr must
# precede shd, tabs, spacing, ind, contextualSpacing, and jc. If we naively
# append pBdr to a paragraph that already has tab stops or spacing set,
# validation fails. _insert_pBdr finds the correct insertion point.

_PBDR_SUCCESSOR_TAGS = ["w:shd", "w:tabs", "w:spacing", "w:ind",
                        "w:contextualSpacing", "w:jc"]


def _insert_pBdr(pPr, pBdr):
    """Insert pBdr into pPr at the correct schema position."""
    successor_qnames = {qn(t) for t in _PBDR_SUCCESSOR_TAGS}
    insert_idx = None
    for i, child in enumerate(pPr):
        if child.tag in successor_qnames:
            insert_idx = i
            break
    if insert_idx is not None:
        pPr.insert(insert_idx, pBdr)
    else:
        pPr.append(pBdr)


def _make_pBdr(*, edge, color, size):
    """Build a w:pBdr element with a single edge border."""
    pBdr = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{edge}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)
    pBdr.append(border)
    return pBdr


def add_paragraph_bottom_border(paragraph, *, color=NAVY, size=8):
    """
    Attach a bottom border to a paragraph. Used for the header strip rule line
    and similar horizontal separators.

    `size` is in eighths of a point per OOXML convention.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = _make_pBdr(edge="bottom", color=color, size=size)
    _insert_pBdr(pPr, pBdr)


def add_paragraph_top_border(paragraph, *, color=BORDER_GREY, size=4):
    """Attach a top border to a paragraph. Used for the footer rule line."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = _make_pBdr(edge="top", color=color, size=size)
    _insert_pBdr(pPr, pBdr)


# --- Helper: cell shading -------------------------------------------------

def set_cell_shading(cell, hex_color):
    """Apply a solid fill color to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# --- Helper: page number field (used inside footer paragraph) -------------

def add_page_number_field(paragraph):
    """
    Insert a Word PAGE field at the current position in `paragraph`.
    Word evaluates the field at open-time to display the current page number.
    """
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# --- Helper: hard page break ---------------------------------------------

def add_page_break(doc):
    """Append a paragraph containing a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


# --- Helper: callout box (single-cell shaded table) -----------------------

def add_callout_box(doc, lines, *, fill=LIGHT_BLUE, bold_first=False, align=None):
    """
    Append a single-cell, full-width table with a shaded fill, used for
    metadata strips and status callouts. `lines` is a list of strings (each
    becomes a paragraph inside the cell). If `bold_first` is True, the first
    line is bolded.
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)

    # Replace default paragraph rather than adding alongside it.
    cell.paragraphs[0].text = ""

    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        if align is not None:
            para.alignment = align
        run = para.add_run(line)
        style_run(run, size=SIZE_BODY, bold=(bold_first and i == 0))

    return table


# --- Helper: dark-navy header callout (Deal Card title block) -------------

def add_navy_header_callout(doc, *, title, lines):
    """
    Single-cell table with NAVY fill and WHITE text. The title line is large
    and bold; subsequent metadata lines are body-size on the same fill.
    Used exclusively for the Deal Card fund-name title block.
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, NAVY)

    cell.paragraphs[0].text = ""
    title_para = cell.paragraphs[0]
    title_run = title_para.add_run(title)
    style_run(title_run, size=SIZE_SUBSECTION_HEADING + 4, bold=True, color=WHITE)

    for line in lines:
        para = cell.add_paragraph()
        run = para.add_run(line)
        style_run(run, size=SIZE_BODY, color=WHITE)

    return table


# --- Helper: section banner (dark-navy filled with white text) ------------

def add_section_banner(doc, text):
    """
    Append a single-cell, full-width table with a dark-navy fill and white
    text — used for major section dividers like 'Section 1 — Market & Focus'.
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, NAVY)

    cell.paragraphs[0].text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    style_run(run, size=SIZE_BANNER, bold=True, color=WHITE)
    return table


# --- Helper: step-icon row (cover page) -----------------------------------

def add_step_icon_row(doc, cells):
    """
    Append a horizontal row of N small cells used on the cover page to
    summarise pipeline stage completion. `cells` is a list of dicts with
    keys: 'present' (bool), 'label_top' (e.g., 'Stage 0'), 'label_bottom'
    (e.g., 'Prescreen').
    """
    n = len(cells)
    table = doc.add_table(rows=1, cols=n)
    table.autofit = False

    for i, c in enumerate(cells):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Replace default paragraph and add three lines: marker, label_top, label_bottom.
        cell.paragraphs[0].text = ""
        marker_para = cell.paragraphs[0]
        marker_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        marker_run = marker_para.add_run("✓" if c["present"] else "—")
        style_run(marker_run, size=14, bold=True, color=NAVY)

        top_para = cell.add_paragraph()
        top_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        top_run = top_para.add_run(c["label_top"])
        style_run(top_run, size=SIZE_BODY, bold=True, color=NAVY)

        bottom_para = cell.add_paragraph()
        bottom_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bottom_run = bottom_para.add_run(c["label_bottom"])
        style_run(bottom_run, size=SIZE_BODY - 1)

    return table


# --- Helper: heading hierarchy ------------------------------------------------

def add_step_heading(doc, text):
    """
    Top-of-step heading: large navy bold all-caps. Used for 'STEP 1 —
    PRESCREEN INTAKE REPORT' and similar major section transitions.
    """
    para = doc.add_paragraph()
    run = para.add_run(text)
    style_run(run, size=SIZE_STEP_HEADING, bold=True, color=NAVY)
    para.paragraph_format.space_after = Pt(12)
    return para


def add_section_heading(doc, text):
    """
    Numbered section heading within a step: '1. Scoring Assessment' style.
    Smaller than step heading, still navy bold.
    """
    para = doc.add_paragraph()
    run = para.add_run(text)
    style_run(run, size=SIZE_SECTION_HEADING, bold=True, color=NAVY)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_subsection_heading(doc, text):
    """
    Sub-section heading: 'Pillar 1 — The Coiled Spring (7/10)' style.
    """
    para = doc.add_paragraph()
    run = para.add_run(text)
    style_run(run, size=SIZE_SUBSECTION_HEADING, bold=True, color=NAVY)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    return para


# --- Helper: prose with bold lead-in -----------------------------------------

def add_lead_in_paragraph(doc, lead_in, body):
    """
    Paragraph where a bold lead-in label is followed by regular body text.
    Used for 'Critical Gap: ...' and 'PCD Intervention Viable?: ...' style
    lines from the legacy 650 report's Section 2 layout.
    """
    para = doc.add_paragraph()
    lead_run = para.add_run(f"{lead_in} ")
    style_run(lead_run, size=SIZE_BODY, bold=True)
    body_run = para.add_run(body)
    style_run(body_run, size=SIZE_BODY)
    return para


# --- Helper: two-column scoring table ----------------------------------------

def add_scoring_table(doc, rows, *, header=("Pillar", "Score")):
    """
    Two-column scoring table with a navy header row.
    `rows` is a list of (label, score_str) tuples. The last row is treated
    as a total row and is bolded.

    Column widths: ~75% label, ~25% score, sized to a 6.5" content area
    (US Letter with 1-inch margins).
    """
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    label_width = Inches(5.0)
    score_width = Inches(1.5)

    # --- Header row ---
    header_cells = table.rows[0].cells
    for i, label in enumerate(header):
        cell = header_cells[i]
        set_cell_shading(cell, NAVY)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(label)
        style_run(run, size=SIZE_BODY, bold=True, color=WHITE)
        cell.width = label_width if i == 0 else score_width

    # --- Data rows ---
    for r, (label, score) in enumerate(rows):
        is_total = (r == len(rows) - 1)
        row_cells = table.rows[r + 1].cells

        # Label cell (left).
        lcell = row_cells[0]
        lcell.paragraphs[0].text = ""
        lrun = lcell.paragraphs[0].add_run(label)
        style_run(lrun, size=SIZE_BODY, bold=is_total)
        lcell.width = label_width

        # Score cell (right).
        scell = row_cells[1]
        scell.paragraphs[0].text = ""
        scell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        srun = scell.paragraphs[0].add_run(score)
        style_run(srun, size=SIZE_BODY, bold=True)
        scell.width = score_width

    return table


# --- Helper: two-column text table ------------------------------------------

def add_two_column_text_table(
    doc, header, rows,
    *, label_width_inches=2.0, body_width_inches=4.5,
):
    """
    Two-column table with text in both columns. Header row navy with white
    text; data rows have bold left column and regular-weight body column,
    both left-aligned. Used for Step 2's Logic & Discipline Element/Finding
    table and similar structures.
    """
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    label_w = Inches(label_width_inches)
    body_w = Inches(body_width_inches)

    # Header row.
    for i, label in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(label)
        style_run(run, size=SIZE_BODY, bold=True, color=WHITE)
        cell.width = label_w if i == 0 else body_w

    # Data rows.
    for r, (label, body) in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        lcell = row_cells[0]
        lcell.paragraphs[0].text = ""
        lrun = lcell.paragraphs[0].add_run(label)
        style_run(lrun, size=SIZE_BODY, bold=True)
        lcell.width = label_w

        bcell = row_cells[1]
        bcell.paragraphs[0].text = ""
        brun = bcell.paragraphs[0].add_run(body)
        style_run(brun, size=SIZE_BODY)
        bcell.width = body_w

    return table


# --- Helper: bullet list ----------------------------------------------------

def add_bullet_list(doc, items):
    """
    Render a list of strings as a bullet-pointed list. Uses Word's built-in
    'List Bullet' paragraph style, which produces standard round-bullet
    formatting that survives faithfully across Word and Pages.
    """
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")
        # Ensure the bulleted paragraph also uses our document font/size.
        for run in para.runs:
            style_run(run, size=SIZE_BODY)
    return None


# --- Helper: multi-column text table (3+ columns) ---------------------------

def add_multi_column_text_table(doc, headers, rows, *, column_widths_inches=None):
    """
    Multi-column table with text in all columns. Header row navy with white
    text; data rows left-aligned, regular weight throughout. Used for Step 4's
    Translation Matrix (5 columns: Theme, Evidence, Primary Match, Secondary
    Matches, Type) and similar wide reference tables.

    `column_widths_inches` is a list of widths, one per column. If None,
    columns are set to equal widths summing to 6.5" (US Letter content area).
    """
    n_cols = len(headers)
    if column_widths_inches is None:
        equal = 6.5 / n_cols
        column_widths_inches = [equal] * n_cols
    assert len(column_widths_inches) == n_cols

    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row.
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(label)
        style_run(run, size=SIZE_BODY - 1, bold=True, color=WHITE)
        cell.width = Inches(column_widths_inches[i])

    # Data rows.
    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.paragraphs[0].text = ""
            text = value if isinstance(value, str) else str(value)
            run = cell.paragraphs[0].add_run(text)
            style_run(run, size=SIZE_BODY - 1)
            cell.width = Inches(column_widths_inches[i])

    return table
