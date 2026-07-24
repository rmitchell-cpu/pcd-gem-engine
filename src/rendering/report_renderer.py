"""
Top-level orchestrator for the pipeline report.

Sets up the Document with US Letter page size, 1-inch margins, default Arial
font, page header (running fund-name strip with rule line), and page footer
(confidential + page number with rule line). Then dispatches each section
renderer in order.

Section renderers are added incrementally across commits:
  Commit 4: cover page only (this file).
  Commit 5: Step 1 — Prescreen Intake Report.
  Commit 6: Step 2 — Deck Analysis Report.
  ... etc.
"""

from datetime import datetime

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt

from .formatters import date_caps
from .io import load_json
from .sections import (
    appendix_evals,
    cover, step1_prescreen, step2_deck_analysis, step3_lp_emails,
    step4_preqin_taxonomy, step5_deal_card,
)
from .styles import (
    BORDER_GREY, FONT, NAVY, SIZE_BODY, SIZE_HEADER_FOOTER,
    add_page_number_field, add_paragraph_bottom_border,
    add_paragraph_top_border, style_run,
)


def _configure_page(doc):
    """US Letter, 1-inch margins on all sides."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def _configure_default_font(doc):
    """Set the document default font so unstyled paragraphs inherit it."""
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SIZE_BODY)


def _configure_header(doc, *, fund_name, as_of_date):
    """
    Page header: 'PCD GP CONCIERGE PIPELINE REPORT — {fund_name}      {DATE}'
    with a navy bottom-border rule line. The fund-name and date are tab-aligned
    so the date sits flush right while the title sits flush left.
    """
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = ""

    # Set a right tab stop so the date lands at the right margin.
    tab_stops = header_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    title_run = header_para.add_run(
        f"PCD GP CONCIERGE PIPELINE REPORT  —  {fund_name}"
    )
    style_run(title_run, size=SIZE_HEADER_FOOTER, bold=True, color=NAVY)

    tab_run = header_para.add_run("\t")
    style_run(tab_run, size=SIZE_HEADER_FOOTER)

    date_run = header_para.add_run(date_caps(as_of_date))
    style_run(date_run, size=SIZE_HEADER_FOOTER, bold=True, color=NAVY)

    add_paragraph_bottom_border(header_para, color=NAVY, size=4)


def _configure_footer(doc):
    """
    Page footer: 'Private Capital Development (PCD) — Confidential' on the
    left, 'Page N' on the right, separated by a tab. A subtle grey top-border
    rule line sits above the footer line.
    """
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""

    add_paragraph_top_border(footer_para, color=BORDER_GREY, size=4)

    tab_stops = footer_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    left_run = footer_para.add_run(
        "Private Capital Development (PCD)  —  Confidential"
    )
    style_run(left_run, size=SIZE_HEADER_FOOTER)

    tab_run = footer_para.add_run("\t")
    style_run(tab_run, size=SIZE_HEADER_FOOTER)

    page_label_run = footer_para.add_run("Page ")
    style_run(page_label_run, size=SIZE_HEADER_FOOTER)
    add_page_number_field(footer_para)


def render_report(artifacts_dir, output_path, presence, *,
                  as_of_date=None, include_evals=True):
    """
    Render the full pipeline report from a job's artifacts directory.

    Args:
        artifacts_dir: Path to jobs/<job_id>/artifacts/
        output_path:   Path where the .docx should be written
        presence:      dict from io.inventory()
        as_of_date:    datetime to display in headers; defaults to today
        include_evals: if True (default), render the eval Appendix when
                       eval_voice.json or eval_cross_stage.json is present.

    Returns the absolute path of the written .docx file.
    """
    if as_of_date is None:
        as_of_date = datetime.now()

    # Load artifacts up-front so the orchestrator can pass them to sections.
    prescreen = load_json(artifacts_dir, "prescreen.json")
    fund_name = (prescreen or {}).get("fund_name", "[Fund Name Missing]")

    # Build the document.
    doc = Document()
    _configure_page(doc)
    _configure_default_font(doc)
    _configure_header(doc, fund_name=fund_name, as_of_date=as_of_date)
    _configure_footer(doc)

    # Section dispatch — pipeline order. Each commit appends another section.
    cover.render(doc, prescreen=prescreen, presence=presence, as_of_date=as_of_date)
    step1_prescreen.render(
        doc, prescreen=prescreen, fund_name=fund_name, as_of_date=as_of_date,
    )
    deck_analysis = load_json(artifacts_dir, "02_deck_analysis.json")
    step2_deck_analysis.render(
        doc, deck_analysis=deck_analysis, fund_name=fund_name,
    )
    lp_emails = load_json(artifacts_dir, "06_lp_emails.json")
    angle_brief = load_json(artifacts_dir, "03_angle_brief.json")
    step3_lp_emails.render(
        doc, lp_emails=lp_emails, angle_brief=angle_brief, fund_name=fund_name,
    )
    preqin_taxonomy = load_json(artifacts_dir, "04_preqin_taxonomy.json")
    step4_preqin_taxonomy.render(doc, preqin_taxonomy=preqin_taxonomy)
    deal_card = load_json(artifacts_dir, "05_deal_card.json")
    step5_deal_card.render(
        doc, deal_card=deal_card, prescreen=prescreen, deck_analysis=deck_analysis,
    )

    if include_evals:
        eval_voice = load_json(artifacts_dir, "eval_voice.json")
        eval_cross_stage = load_json(artifacts_dir, "eval_cross_stage.json")
        appendix_evals.render(
            doc, eval_voice=eval_voice, eval_cross_stage=eval_cross_stage,
        )

    # Save.
    output_path = output_path.expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
