"""
Cover page renderer.

Produces the first page of the report:
  - Header strip (handled at section level by report_renderer)
  - Big navy title 'GP CONCIERGE PIPELINE REPORT'
  - Orange subtitle (fund name)
  - Horizontal rule
  - 'Prepared by Private Capital Development (PCD)' + date
  - PIPELINE STATUS callout (light-blue) with prescreen score and tier
  - Step-icon row showing which stages produced artifacts
  - Hard page break to land Step 1 on its own page
"""

from datetime import datetime

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..formatters import date_mixed, tier_display
from ..styles import (
    NAVY, ORANGE,
    SIZE_COVER_TITLE, SIZE_COVER_SUBTITLE,
    add_callout_box, add_page_break, add_paragraph_bottom_border,
    add_step_icon_row, style_run,
)


# Stages shown in the cover-page step-icon row.
# (filename, label_top, label_bottom). The row is presence-aware, so jobs
# run before 01_fund_extract was implemented (24 July 2026) simply show
# Stage 1 as not produced.
COVER_STAGES = [
    ("prescreen.json",          "Stage 0", "Prescreen"),
    ("01_fund_extract.json",    "Stage 1", "Fund Extract"),
    ("02_deck_analysis.json",   "Stage 2", "Deck Analysis"),
    ("03_angle_brief.json",     "Stage 3", "Angle Brief"),
    ("04_preqin_taxonomy.json", "Stage 4", "Preqin Taxonomy"),
    ("05_deal_card.json",       "Stage 5", "Deal Card"),
    ("06_lp_emails.json",       "Stage 6", "LP Emails"),
]


def render(doc, *, prescreen, presence, as_of_date=None):
    """
    Render the full cover page into `doc`.

    Args:
        doc:         the python-docx Document being built
        prescreen:   parsed prescreen.json (or None)
        presence:    dict from io.inventory(), used for stage-icon checkmarks
        as_of_date:  datetime to display; defaults to today
    """
    if as_of_date is None:
        as_of_date = datetime.now()

    fund_name = (prescreen or {}).get("fund_name", "[Fund Name Missing]")
    total_score = (prescreen or {}).get("total_score", "?")
    classification = (prescreen or {}).get("classification", "")

    # --- Big title ----------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Vertical breathing room above the title.
    title_para.paragraph_format.space_before = None
    title_run = title_para.add_run("GP CONCIERGE PIPELINE REPORT")
    style_run(title_run, size=SIZE_COVER_TITLE, bold=True, color=NAVY)

    # --- Orange subtitle (fund name) ----------------------------------------
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(fund_name)
    style_run(subtitle_run, size=SIZE_COVER_SUBTITLE, bold=True, color=ORANGE)

    # --- Horizontal rule below the title block ------------------------------
    rule_para = doc.add_paragraph()
    add_paragraph_bottom_border(rule_para, color=NAVY, size=8)

    # --- Prepared-by + date -------------------------------------------------
    prepared_para = doc.add_paragraph()
    prepared_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(prepared_para.add_run("Prepared by Private Capital Development (PCD)"))

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(date_para.add_run(date_mixed(as_of_date)))

    # Spacer.
    doc.add_paragraph()

    # --- PIPELINE STATUS callout --------------------------------------------
    status_lines = [
        "PIPELINE STATUS",
        f"Prescreen Score: {total_score}/40",
        f"Classification: {tier_display(classification)}",
    ]
    add_callout_box(
        doc,
        status_lines,
        bold_first=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Spacer.
    doc.add_paragraph()

    # --- Step-icon row ------------------------------------------------------
    cells = [
        {
            "present": presence.get(filename, False),
            "label_top": top,
            "label_bottom": bottom,
        }
        for (filename, top, bottom) in COVER_STAGES
    ]
    add_step_icon_row(doc, cells)

    # --- Page break to Step 1 ----------------------------------------------
    add_page_break(doc)
